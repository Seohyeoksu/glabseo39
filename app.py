import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
from datetime import datetime, timedelta
import calendar

from docx.enum.section import WD_SECTION

def add_footer(doc):
    """페이지 하단에 푸터 추가"""
    # 모든 섹션에 푸터 추가
    for section in doc.sections:
        footer = section.footer
        
        # 푸터가 비어있으면 새 단락 추가
        if not footer.paragraphs:
            footer_para = footer.add_paragraph()
        else:
            footer_para = footer.paragraphs[0]
            footer_para.clear()
        
        # 푸터 텍스트 추가
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run("세계교육 표준으로 삶의 힘을 키우는 따뜻한 경북교육")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer_run.font.italic = True
        
        # 여백 조정
        footer_para.paragraph_format.space_before = Pt(12)

def add_user_info(doc, school_name="", grade="", class_num="", student_name=""):
    """페이지 상단에 사용자 정보 추가"""
    if any([school_name, grade, class_num, student_name]):
        # 사용자 정보 테이블
        info_table = doc.add_table(rows=1, cols=4)
        info_table.style = 'Normal Table'
        info_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        
        # 학교명
        if school_name:
            school_cell = info_table.cell(0, 0)
            school_cell.width = Inches(2)
            school_p = school_cell.paragraphs[0]
            school_p.add_run(school_name)
            school_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 학년
        if grade:
            grade_cell = info_table.cell(0, 1)
            grade_cell.width = Inches(1)
            grade_p = grade_cell.paragraphs[0]
            grade_p.add_run(grade)
            grade_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 반
        if class_num:
            class_cell = info_table.cell(0, 2)
            class_cell.width = Inches(1)
            class_p = class_cell.paragraphs[0]
            class_p.add_run(class_num)
            class_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 이름
        if student_name:
            name_cell = info_table.cell(0, 3)
            name_cell.width = Inches(1.5)
            name_p = name_cell.paragraphs[0]
            name_p.add_run(f"이름: {student_name}")
            name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 테이블 스타일 조정
        for row in info_table.rows:
            for cell in row.cells:
                # 테두리 제거
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border in ['top', 'left', 'bottom', 'right']:
                    side = OxmlElement(f'w:{border}')
                    side.set(qn('w:val'), 'nil')
                    tcBorders.append(side)
                tcPr.append(tcBorders)
                
                # 폰트 크기 조정
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        
        # 구분선
        line_para = doc.add_paragraph("─" * 80)
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_para.paragraph_format.space_before = Pt(6)
        line_para.paragraph_format.space_after = Pt(12)
        
        return True
    return False

def create_lined_notebook(doc, lines_per_page=25, num_pages=5, user_info=None):
    """줄공책 양식 생성 - 테이블 방식"""
    for page in range(num_pages):
        if page > 0:
            doc.add_page_break()
        
        # 사용자 정보 추가
        if user_info and page == 0:
            add_user_info(doc, **user_info)
        
        # 페이지 상단 여백
        top_para = doc.add_paragraph()
        top_para.paragraph_format.space_after = Pt(10)
        
        # 테이블을 사용한 줄 생성
        table = doc.add_table(rows=lines_per_page, cols=1)
        table.autofit = False
        table.style = 'Normal Table'
        
        for i, row in enumerate(table.rows):
            # 행 높이 설정
            row.height = Pt(28)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            
            cell = row.cells[0]
            cell.width = Inches(7.5)
            
            # 셀 내부 단락 설정
            if cell.paragraphs:
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            
            # 셀 테두리 설정
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            
            # 기존 테두리 제거
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is not None:
                tcPr.remove(tcBorders)
            
            # 새 테두리 설정
            tcBorders = OxmlElement('w:tcBorders')
            
            # 하단 선만 추가
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), '808080')
            tcBorders.append(bottom)
            
            # 나머지 테두리는 없음
            for border in ['top', 'left', 'right']:
                side = OxmlElement(f'w:{border}')
                side.set(qn('w:val'), 'nil')
                tcBorders.append(side)
            
            tcPr.append(tcBorders)
            
            # 셀 여백 설정
            tcMar = tcPr.find(qn('w:tcMar'))
            if tcMar is not None:
                tcPr.remove(tcMar)
                
            tcMar = OxmlElement('w:tcMar')
            for margin_name in ['top', 'left', 'bottom', 'right']:
                margin = OxmlElement(f'w:{margin_name}')
                margin.set(qn('w:w'), '50')
                margin.set(qn('w:type'), 'dxa')
                tcMar.append(margin)
            tcPr.append(tcMar)

def create_grid_notebook(doc, rows=15, cols=15, num_pages=5, user_info=None):
    """칸공책 양식 생성"""
    # 첫 페이지에서 푸터 설정
    add_footer(doc)
    
    for page in range(num_pages):
        if page > 0:
            doc.add_page_break()
        
        # 사용자 정보 추가
        if user_info and page == 0:
            add_user_info(doc, **user_info)
        
        # 페이지 크기 계산 (A4 기준)
        page_width = 8.27 - 1.0  # 인치 (여백 제외)
        page_height = 11.69 - 1.0
        
        cell_width = page_width / cols
        cell_height = page_height / rows
        
        # 테이블 생성
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.allow_autofit = False
        
        # 각 행 설정
        for row in table.rows:
            # 행 높이 설정
            tr = row._element
            trPr = tr.get_or_add_trPr()
            
            # 기존 높이 설정 제거
            for child in trPr:
                if child.tag.endswith('trHeight'):
                    trPr.remove(child)
            
            # 새 높이 설정
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(int(cell_height * 1440)))  # twips
            trHeight.set(qn('w:hRule'), 'exact')
            trPr.append(trHeight)
            
            # 각 셀 설정
            for cell in row.cells:
                # 셀 너비 설정
                cell.width = Inches(cell_width)
                
                # 셀 내용 설정
                if cell.paragraphs:
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = Pt(0)
                
                # 셀 여백 최소화
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                
                # 기존 여백 제거
                tcMar = tcPr.find(qn('w:tcMar'))
                if tcMar is not None:
                    tcPr.remove(tcMar)
                
                # 새 여백 설정
                tcMar = OxmlElement('w:tcMar')
                for margin_name in ['top', 'left', 'bottom', 'right']:
                    margin = OxmlElement(f'w:{margin_name}')
                    margin.set(qn('w:w'), '10')
                    margin.set(qn('w:type'), 'dxa')
                    tcMar.append(margin)
                tcPr.append(tcMar)

def create_english_notebook(doc, lines_per_page=12, num_pages=5, user_info=None):
    """영어노트 양식 생성 (4선 노트)"""
    # 첫 페이지에서 푸터 설정
    add_footer(doc)
    
    for page in range(num_pages):
        if page > 0:
            doc.add_page_break()
        
        # 사용자 정보 추가
        if user_info and page == 0:
            add_user_info(doc, **user_info)
        
        # 페이지 상단 여백
        top_margin = doc.add_paragraph()
        top_margin.paragraph_format.space_after = Pt(20)
        
        # 페이지 높이 계산 (A4 기준, 여백 제외)
        available_height = 11.69 - 1.0  # 인치
        header_space = 1.5  # 상단 여백
        remaining_height = available_height - header_space
        
        # 줄 간격 계산 (줄 수에 따라 동적으로 조정)
        total_spacing = remaining_height / lines_per_page
        line_spacing = total_spacing * 0.8  # 80%는 줄 간격
        between_spacing = total_spacing * 0.2  # 20%는 줄 사이 간격
        
        for i in range(lines_per_page):
            # 4선을 위한 테이블 생성
            table = doc.add_table(rows=4, cols=1)
            table.autofit = False
            table.style = 'Normal Table'
            
            # 각 선의 높이 비율
            line_heights = [
                line_spacing * 0.2,  # 상단 점선
                line_spacing * 0.2,  # 상단 실선
                line_spacing * 0.3,  # 기준선 (더 넓게)
                line_spacing * 0.3   # 하단 실선
            ]
            
            # 첫 번째 선 (상단 점선)
            row1 = table.rows[0]
            row1.height = Pt(line_heights[0] * 72)  # 인치를 포인트로 변환
            row1.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            cell1 = row1.cells[0]
            cell1.width = Inches(7.5)
            
            # 점선 스타일
            tc1 = cell1._element
            tcPr1 = tc1.get_or_add_tcPr()
            tcBorders1 = OxmlElement('w:tcBorders')
            bottom1 = OxmlElement('w:bottom')
            bottom1.set(qn('w:val'), 'dotted')
            bottom1.set(qn('w:sz'), '4')
            bottom1.set(qn('w:color'), 'CCCCCC')
            tcBorders1.append(bottom1)
            tcPr1.append(tcBorders1)
            
            # 두 번째 선 (상단 실선)
            row2 = table.rows[1]
            row2.height = Pt(line_heights[1] * 72)
            row2.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            cell2 = row2.cells[0]
            
            tc2 = cell2._element
            tcPr2 = tc2.get_or_add_tcPr()
            tcBorders2 = OxmlElement('w:tcBorders')
            bottom2 = OxmlElement('w:bottom')
            bottom2.set(qn('w:val'), 'single')
            bottom2.set(qn('w:sz'), '4')
            bottom2.set(qn('w:color'), '808080')
            tcBorders2.append(bottom2)
            tcPr2.append(tcBorders2)
            
            # 세 번째 선 (기준선 - 굵은 실선)
            row3 = table.rows[2]
            row3.height = Pt(line_heights[2] * 72)
            row3.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            cell3 = row3.cells[0]
            
            tc3 = cell3._element
            tcPr3 = tc3.get_or_add_tcPr()
            tcBorders3 = OxmlElement('w:tcBorders')
            bottom3 = OxmlElement('w:bottom')
            bottom3.set(qn('w:val'), 'single')
            bottom3.set(qn('w:sz'), '6')
            bottom3.set(qn('w:color'), '000000')
            tcBorders3.append(bottom3)
            tcPr3.append(tcBorders3)
            
            # 네 번째 선 (하단 실선)
            row4 = table.rows[3]
            row4.height = Pt(line_heights[3] * 72)
            row4.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            cell4 = row4.cells[0]
            
            tc4 = cell4._element
            tcPr4 = tc4.get_or_add_tcPr()
            tcBorders4 = OxmlElement('w:tcBorders')
            bottom4 = OxmlElement('w:bottom')
            bottom4.set(qn('w:val'), 'single')
            bottom4.set(qn('w:sz'), '4')
            bottom4.set(qn('w:color'), '808080')
            tcBorders4.append(bottom4)
            tcPr4.append(tcBorders4)
            
            # 모든 셀의 다른 테두리 제거
            for row in table.rows:
                cell = row.cells[0]
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is not None:
                    for border in ['top', 'left', 'right']:
                        side = tcBorders.find(qn(f'w:{border}'))
                        if side is None:
                            side = OxmlElement(f'w:{border}')
                            side.set(qn('w:val'), 'nil')
                            tcBorders.append(side)
            
            # 줄 사이 간격 (마지막 줄 제외)
            if i < lines_per_page - 1:
                spacing = doc.add_paragraph()
                spacing.paragraph_format.space_after = Pt(between_spacing * 72)

def create_cornell_notebook(doc, num_pages=5, user_info=None):
    """코넬노트 양식 생성"""
    # 첫 페이지에서 푸터 설정
    add_footer(doc)
    
    for page in range(num_pages):
        if page > 0:
            doc.add_page_break()
        
        # 사용자 정보 추가
        if user_info and page == 0:
            add_user_info(doc, **user_info)
        
        # 상단 영역 (제목, 날짜)
        header_table = doc.add_table(rows=1, cols=2)
        header_table.style = 'Table Grid'
        header_table.columns[0].width = Inches(4)
        header_table.columns[1].width = Inches(2.5)
        
        # 제목 셀
        title_cell = header_table.cell(0, 0)
        title_p = title_cell.paragraphs[0]
        title_p.add_run("제목: ").bold = True
        
        # 날짜 셀
        date_cell = header_table.cell(0, 1)
        date_p = date_cell.paragraphs[0]
        date_p.add_run("날짜: ").bold = True
        
        # 간격
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
        
        # 메인 영역 (핵심어 | 노트)
        main_table = doc.add_table(rows=1, cols=2)
        main_table.style = 'Table Grid'
        main_table.columns[0].width = Inches(2)
        main_table.columns[1].width = Inches(4.5)
        
        # 핵심어 열
        key_cell = main_table.cell(0, 0)
        key_p = key_cell.paragraphs[0]
        key_p.add_run("핵심어/질문").bold = True
        key_p.add_run("\n\n")
        
        # 노트 열
        note_cell = main_table.cell(0, 1)
        note_p = note_cell.paragraphs[0]
        note_p.add_run("노트 영역").bold = True
        note_p.add_run("\n\n")
        
        # 셀 높이 설정
        tr = main_table.rows[0]._element
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '8000')  # 약 5.5인치
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)
        
        # 간격
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
        
        # 하단 요약 영역
        summary_title = doc.add_paragraph("요약:")
        summary_title.runs[0].font.bold = True
        summary_title.paragraph_format.space_after = Pt(6)
        
        # 요약 박스
        summary_table = doc.add_table(rows=1, cols=1)
        summary_table.style = 'Table Grid'
        summary_cell = summary_table.cell(0, 0)
        
        # 요약 영역 높이 설정
        tr = summary_table.rows[0]._element
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '2000')  # 약 1.5인치
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)

def create_music_staff(doc, staves_per_page=12, num_pages=5, user_info=None):
    """음악 오선지 생성"""
    # 첫 페이지에서 푸터 설정
    add_footer(doc)
    
    for page in range(num_pages):
        if page > 0:
            doc.add_page_break()
        
        # 사용자 정보 추가
        if user_info and page == 0:
            add_user_info(doc, **user_info)
        
        # 페이지 상단 여백
        top_para = doc.add_paragraph()
        top_para.paragraph_format.space_after = Pt(20)
        
        # 페이지 높이 계산 (A4 기준, 여백 제외)
        available_height = 11.69 - 1.0  # 인치
        header_space = 1.5  # 상단 여백
        remaining_height = available_height - header_space
        
        # 오선지당 높이 계산
        staff_total_height = remaining_height / staves_per_page
        staff_height = staff_total_height * 0.4  # 40%는 오선지
        spacing_height = staff_total_height * 0.6  # 60%는 간격
        
        for staff_num in range(staves_per_page):
            # 5선을 위한 테이블 생성
            table = doc.add_table(rows=5, cols=1)
            table.autofit = False
            table.style = 'Normal Table'
            
            # 각 선의 간격 계산
            line_spacing = staff_height / 5
            
            for i, row in enumerate(table.rows):
                # 행 높이 설정
                row.height = Pt(line_spacing * 72)  # 인치를 포인트로 변환
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                
                cell = row.cells[0]
                cell.width = Inches(7.5)
                
                # 셀 테두리 설정 (하단 선만)
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                
                # 기존 테두리 제거
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is not None:
                    tcPr.remove(tcBorders)
                
                # 새 테두리 설정
                tcBorders = OxmlElement('w:tcBorders')
                
                # 하단 선만 추가
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:space'), '0')
                bottom.set(qn('w:color'), '000000')
                tcBorders.append(bottom)
                
                # 나머지 테두리는 없음
                for border in ['top', 'left', 'right']:
                    side = OxmlElement(f'w:{border}')
                    side.set(qn('w:val'), 'nil')
                    tcBorders.append(side)
                
                tcPr.append(tcBorders)
                
                # 셀 여백 제거
                tcMar = OxmlElement('w:tcMar')
                for margin_name in ['top', 'left', 'bottom', 'right']:
                    margin = OxmlElement(f'w:{margin_name}')
                    margin.set(qn('w:w'), '0')
                    margin.set(qn('w:type'), 'dxa')
                    tcMar.append(margin)
                tcPr.append(tcMar)
            
            # 오선 사이 간격 (마지막 오선 제외)
            if staff_num < staves_per_page - 1:
                spacing = doc.add_paragraph()
                spacing.paragraph_format.space_after = Pt(spacing_height * 72)

def create_chinese_notebook(doc, rows_per_page=6, chars_per_row=8, num_pages=5, user_info=None):
    """한자 노트 생성 - 한국식 한자 쓰기 노트"""
    # 첫 페이지에서 푸터 설정
    add_footer(doc)
    
    for page in range(num_pages):
        if page > 0:
            doc.add_page_break()
        
        # 사용자 정보 추가
        if user_info and page == 0:
            add_user_info(doc, **user_info)
        
        # 페이지 상단 여백
        top_para = doc.add_paragraph()
        top_para.paragraph_format.space_after = Pt(20)
        
        # 페이지 높이 계산 (A4 기준, 여백 제외)
        available_height = 11.69 - 1.0  # 인치
        header_space = 1.5  # 상단 여백
        remaining_height = available_height - header_space
        
        # 행당 높이 계산
        row_total_height = remaining_height / rows_per_page
        hanja_cell_height = row_total_height * 0.7  # 70%는 한자 칸
        meaning_cell_height = row_total_height * 0.2  # 20%는 뜻 칸
        spacing_height = row_total_height * 0.1  # 10%는 간격
        
        # 한자 연습용 테이블 생성 (한자칸 + 뜻칸)
        for row_idx in range(rows_per_page):
            # 한 줄에 한자칸과 뜻칸을 함께 생성
            line_table = doc.add_table(rows=2, cols=chars_per_row)
            line_table.style = 'Table Grid'
            line_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            line_table.autofit = False
            
            # 첫 번째 행: 한자 쓰기 칸
            hanja_row = line_table.rows[0]
            hanja_row.height = Pt(hanja_cell_height * 72)  # 인치를 포인트로 변환
            hanja_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            
            for col_idx in range(chars_per_row):
                cell = hanja_row.cells[col_idx]
                cell.width = Pt(hanja_cell_height * 72)  # 정사각형으로 만들기
                
                # 십자 가이드라인을 위한 2x2 내부 테이블
                guide_table = cell.add_table(rows=2, cols=2)
                guide_table.autofit = False
                
                # 4개의 셀로 십자 만들기
                for i in range(2):
                    for j in range(2):
                        guide_cell = guide_table.cell(i, j)
                        
                        # 열 너비 설정 (왼쪽을 살짝 좁게)
                        if j == 0:
                            guide_cell.width = Pt(hanja_cell_height * 72 * 0.45)  # 왼쪽 열: 살짝 좁게
                        else:
                            guide_cell.width = Pt(hanja_cell_height * 72 * 0.55)  # 오른쪽 열: 살짝 넓게
                        
                        tc = guide_cell._element
                        tcPr = tc.get_or_add_tcPr()
                        
                        # 행 높이 설정 (위쪽 행을 살짝 작게)
                        if i == 0:
                            tcH = OxmlElement('w:tcH')
                            tcH.set(qn('w:val'), str(int(hanja_cell_height * 72 * 0.45)))  # 위쪽 행: 살짝 작게
                            tcH.set(qn('w:hRule'), 'exact')
                            tcPr.append(tcH)
                        else:
                            tcH = OxmlElement('w:tcH')
                            tcH.set(qn('w:val'), str(int(hanja_cell_height * 72 * 0.55)))  # 아래쪽 행: 살짝 크게
                            tcH.set(qn('w:hRule'), 'exact')
                            tcPr.append(tcH)
                        
                        # 테두리 설정 - 내부 선만 점선으로
                        tcBorders = OxmlElement('w:tcBorders')
                        
                        # 왼쪽 위 셀
                        if i == 0 and j == 0:
                            borders = {'right': 'dotted', 'bottom': 'dotted'}
                        # 오른쪽 위 셀
                        elif i == 0 and j == 1:
                            borders = {'left': 'dotted', 'bottom': 'dotted'}
                        # 왼쪽 아래 셀
                        elif i == 1 and j == 0:
                            borders = {'right': 'dotted', 'top': 'dotted'}
                        # 오른쪽 아래 셀
                        else:
                            borders = {'left': 'dotted', 'top': 'dotted'}
                        
                        # 점선 테두리 추가
                        for border, style in borders.items():
                            side = OxmlElement(f'w:{border}')
                            side.set(qn('w:val'), style)
                            side.set(qn('w:sz'), '6')
                            side.set(qn('w:color'), 'CCCCCC')
                            tcBorders.append(side)
                        
                        # 외곽선은 없음
                        for border in ['top', 'bottom', 'left', 'right']:
                            if border not in borders:
                                side = OxmlElement(f'w:{border}')
                                side.set(qn('w:val'), 'nil')
                                tcBorders.append(side)
                        
                        tcPr.append(tcBorders)
                        
                        # 여백 제거
                        tcMar = OxmlElement('w:tcMar')
                        for margin in ['top', 'left', 'bottom', 'right']:
                            m = OxmlElement(f'w:{margin}')
                            m.set(qn('w:w'), '0')
                            m.set(qn('w:type'), 'dxa')
                            tcMar.append(m)
                        tcPr.append(tcMar)
            
            # 두 번째 행: 뜻 쓰기 칸
            meaning_row = line_table.rows[1]
            meaning_row.height = Pt(meaning_cell_height * 72)
            meaning_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            
            for col_idx in range(chars_per_row):
                cell = meaning_row.cells[col_idx]
                cell.width = Pt(hanja_cell_height * 72)
                
                # 뜻 칸 스타일
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                
                # 연한 배경색
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                shading = OxmlElement('w:shd')
                shading.set(qn('w:val'), 'clear')
                shading.set(qn('w:color'), 'auto')
                shading.set(qn('w:fill'), 'F5F5F5')
                tcPr.append(shading)
            
            # 줄 간격 (마지막 줄 제외)
            if row_idx < rows_per_page - 1:
                spacing = doc.add_paragraph()
                spacing.paragraph_format.space_after = Pt(spacing_height * 72)

def create_diary(doc, start_date, num_days, user_info=None):
    """다이어리 양식 생성"""
    # 첫 페이지에서 푸터 설정
    add_footer(doc)
    
    # 첫 페이지에 사용자 정보 추가
    if user_info:
        add_user_info(doc, **user_info)
        doc.add_page_break()
    
    for day in range(num_days):
        if day > 0:
            doc.add_page_break()
        
        current_date = start_date + timedelta(days=day)
        
        # 날짜 헤더
        date_header = doc.add_paragraph()
        date_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_header.add_run(current_date.strftime("%Y년 %m월 %d일 %A"))
        date_run.font.size = Pt(16)
        date_run.font.bold = True
        
        # 날씨, 기분, 중요도
        info_table = doc.add_table(rows=1, cols=3)
        info_table.style = 'Light List'
        
        weather_cell = info_table.cell(0, 0)
        weather_cell.text = "날씨: ☀️ ☁️ 🌧️ ❄️"
        
        mood_cell = info_table.cell(0, 1)
        mood_cell.text = "기분: 😊 😐 😢 😡"
        
        importance_cell = info_table.cell(0, 2)
        importance_cell.text = "중요도: ⭐⭐⭐⭐⭐"
        
        # 간격
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
        
        # 일정 표
        schedule_title = doc.add_paragraph("📅 오늘의 일정")
        schedule_title.runs[0].font.bold = True
        
        schedule_table = doc.add_table(rows=10, cols=2)
        schedule_table.style = 'Light Grid'
        
        # 시간대별 일정
        times = ["오전 7-9시", "오전 9-11시", "오전 11시-오후 1시", 
                "오후 1-3시", "오후 3-5시", "오후 5-7시", "오후 7-9시", 
                "오후 9-11시", "기타", "메모"]
        
        for i, time in enumerate(times):
            time_cell = schedule_table.cell(i, 0)
            time_cell.text = time
            time_cell.width = Inches(1.5)
            
            content_cell = schedule_table.cell(i, 1)
            content_cell.width = Inches(5)
        
        # 간격
        doc.add_paragraph().paragraph_format.space_after = Pt(20)
        
        # 일기 작성 공간
        diary_title = doc.add_paragraph("✍️ 오늘의 일기")
        diary_title.runs[0].font.bold = True
        diary_title.paragraph_format.space_after = Pt(12)
        
        # 줄 노트 추가
        diary_table = doc.add_table(rows=15, cols=1)
        diary_table.style = 'Normal Table'
        
        for row in diary_table.rows:
            row.height = Pt(25)
            cell = row.cells[0]
            
            # 하단 선 추가
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '2')
            bottom.set(qn('w:color'), 'CCCCCC')
            tcBorders.append(bottom)
            
            for border in ['top', 'left', 'right']:
                side = OxmlElement(f'w:{border}')
                side.set(qn('w:val'), 'nil')
                tcBorders.append(side)
            
            tcPr.append(tcBorders)
        
        # 감사 일기
        doc.add_paragraph()
        gratitude_title = doc.add_paragraph("🙏 오늘 감사한 일 3가지")
        gratitude_title.runs[0].font.bold = True
        
        for i in range(3):
            gratitude = doc.add_paragraph(f"{i+1}. ", style='List Number')
            gratitude.paragraph_format.space_after = Pt(12)
        
def create_calendar(doc, year, month, num_months=12, user_info=None):
    """달력 양식 생성"""
    # 첫 페이지에서 푸터 설정
    add_footer(doc)
    
    # 첫 페이지에 사용자 정보 추가
    if user_info:
        add_user_info(doc, **user_info)
        doc.add_page_break()
    
    for i in range(num_months):
        if i > 0:
            doc.add_page_break()
        
        current_month = month + i
        current_year = year
        if current_month > 12:
            current_year += (current_month - 1) // 12
            current_month = ((current_month - 1) % 12) + 1
        
        # 월 제목
        month_names = ['', '1월', '2월', '3월', '4월', '5월', '6월', 
                      '7월', '8월', '9월', '10월', '11월', '12월']
        title = doc.add_paragraph()
        title_run = title.add_run(f"{current_year}년 {month_names[current_month]}")
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(12)
        
        # 요일 헤더
        weekdays = ['월', '화', '수', '목', '금', '토', '일']
        
        # 달력 테이블 생성 (요일 + 최대 6주)
        table = doc.add_table(rows=7, cols=7)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 요일 헤더 설정
        header_row = table.rows[0]
        for j, day in enumerate(weekdays):
            cell = header_row.cells[j]
            cell.text = day
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.bold = True
            
            # 토요일은 파란색, 일요일은 빨간색
            if j == 5:  # 토요일
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 255)
            elif j == 6:  # 일요일
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
        
        # 달력 날짜 채우기
        cal = calendar.monthcalendar(current_year, current_month)
        
        for week_num, week in enumerate(cal):
            row = table.rows[week_num + 1]
            for day_num, day in enumerate(week):
                cell = row.cells[day_num]
                
                if day != 0:
                    # 날짜 추가
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    date_run = p.add_run(str(day))
                    date_run.font.size = Pt(10)
                    date_run.font.bold = True
                    
                    # 주말 색상
                    if day_num == 5:  # 토요일
                        date_run.font.color.rgb = RGBColor(0, 0, 255)
                    elif day_num == 6:  # 일요일
                        date_run.font.color.rgb = RGBColor(255, 0, 0)
                    
                    # 메모 공간을 위한 줄바꿈
                    p.add_run('\n\n\n')
                
                # 셀 크기 설정
                cell.width = Inches(1)
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcH = OxmlElement('w:tcH')
                tcH.set(qn('w:val'), '1500')
                tcH.set(qn('w:hRule'), 'atLeast')
                tcPr.append(tcH)
        
        # 하단 메모 영역
        doc.add_paragraph()
        memo_title = doc.add_paragraph("📝 이달의 메모")
        memo_title.runs[0].font.bold = True
        memo_title.paragraph_format.space_after = Pt(6)
        
        memo_table = doc.add_table(rows=3, cols=1)
        memo_table.style = 'Light List'
        for row in memo_table.rows:
            row.height = Pt(40)

def create_math_error_notebook(doc, problems_per_page=3, num_pages=5, user_info=None):
    """수학 오답 노트 생성"""
    # 첫 페이지에서 푸터 설정
    add_footer(doc)
    
    for page in range(num_pages):
        if page > 0:
            doc.add_page_break()
        
        # 사용자 정보 추가
        if user_info and page == 0:
            add_user_info(doc, **user_info)
        
        # 페이지 헤더
        header = doc.add_paragraph()
        header_run = header.add_run(f"수학 오답 노트 - {page + 1}페이지")
        header_run.font.size = Pt(14)
        header_run.font.bold = True
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.paragraph_format.space_after = Pt(20)
        
        # 페이지 높이 계산 (A4 기준, 여백 제외)
        available_height = 11.69 - 1.0  # 인치
        header_space = 1.5  # 헤더와 여백
        remaining_height = available_height - header_space
        
        # 문제당 할당 높이 계산
        section_height = remaining_height / problems_per_page
        
        # 각 섹션의 구성 요소별 높이 비율
        info_height = section_height * 0.08
        prob_height = section_height * 0.25
        solution_height = section_height * 0.45
        analysis_height = section_height * 0.17
        separator_height = section_height * 0.05
        
        # 각 문제별 섹션
        for prob_num in range(problems_per_page):
            # 문제 정보 테이블
            info_table = doc.add_table(rows=1, cols=4)
            info_table.style = 'Table Grid'
            
            # 문제 번호
            prob_cell = info_table.cell(0, 0)
            prob_cell.width = Inches(1.5)
            prob_p = prob_cell.paragraphs[0]
            prob_p.add_run("문제 번호:").bold = True
            
            # 날짜
            date_cell = info_table.cell(0, 1)
            date_cell.width = Inches(1.5)
            date_p = date_cell.paragraphs[0]
            date_p.add_run("날짜:").bold = True
            
            # 출처
            source_cell = info_table.cell(0, 2)
            source_cell.width = Inches(2)
            source_p = source_cell.paragraphs[0]
            source_p.add_run("출처:").bold = True
            
            # 난이도
            level_cell = info_table.cell(0, 3)
            level_cell.width = Inches(1.5)
            level_p = level_cell.paragraphs[0]
            level_p.add_run("난이도: ⭐⭐⭐⭐⭐").bold = True
            
            # 간격
            spacing1 = doc.add_paragraph()
            spacing1.paragraph_format.space_after = Pt(6)
            
            # 문제 영역
            prob_title = doc.add_paragraph("📝 문제")
            prob_title.runs[0].font.bold = True
            prob_title.runs[0].font.size = Pt(11)
            prob_title.paragraph_format.space_after = Pt(4)
            
            prob_table = doc.add_table(rows=1, cols=1)
            prob_table.style = 'Table Grid'
            prob_content_cell = prob_table.cell(0, 0)
            
            # 문제 영역 높이 설정 (문제 수에 따라 조정)
            tc = prob_content_cell._element
            tcPr = tc.get_or_add_tcPr()
            tcH = OxmlElement('w:tcH')
            tcH.set(qn('w:val'), str(int(prob_height * 1440)))  # twips 변환
            tcH.set(qn('w:hRule'), 'exact')
            tcPr.append(tcH)
            
            # 문제 영역 배경색
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), 'F0F8FF')
            tcPr.append(shading)
            
            # 간격
            spacing2 = doc.add_paragraph()
            spacing2.paragraph_format.space_after = Pt(6)
            
            # 풀이 과정 영역
            solution_title = doc.add_paragraph("✏️ 풀이 과정")
            solution_title.runs[0].font.bold = True
            solution_title.runs[0].font.size = Pt(11)
            solution_title.paragraph_format.space_after = Pt(4)
            
            # 격자 노트 스타일 테이블 (문제 수에 따라 행 수 조정)
            grid_rows = max(6, int(20 / problems_per_page))
            solution_table = doc.add_table(rows=grid_rows, cols=15)
            solution_table.style = 'Table Grid'
            solution_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # 각 행의 높이를 문제 수에 따라 조정
            row_height = int(solution_height * 1440 / grid_rows)
            
            for row in solution_table.rows:
                # 행 높이 설정
                tr = row._element
                trPr = tr.get_or_add_trPr()
                trHeight = OxmlElement('w:trHeight')
                trHeight.set(qn('w:val'), str(row_height))
                trHeight.set(qn('w:hRule'), 'exact')
                trPr.append(trHeight)
                
                for cell in row.cells:
                    cell.width = Pt(30)
                    
                    # 연한 격자선
                    tc = cell._element
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = OxmlElement('w:tcBorders')
                    
                    for border in ['top', 'left', 'bottom', 'right']:
                        side = OxmlElement(f'w:{border}')
                        side.set(qn('w:val'), 'single')
                        side.set(qn('w:sz'), '2')
                        side.set(qn('w:color'), 'E0E0E0')
                        tcBorders.append(side)
                    
                    tcPr.append(tcBorders)
            
            # 간격
            spacing3 = doc.add_paragraph()
            spacing3.paragraph_format.space_after = Pt(6)
            
            # 오답 원인 및 핵심 포인트
            analysis_table = doc.add_table(rows=1, cols=2)
            analysis_table.style = 'Table Grid'
            
            # 오답 원인
            cause_cell = analysis_table.cell(0, 0)
            cause_cell.width = Inches(3.25)
            cause_p = cause_cell.paragraphs[0]
            cause_p.add_run("❌ 오답 원인").bold = True
            cause_p.add_run("\n\n□ 개념 이해 부족\n□ 계산 실수\n□ 문제 해석 오류\n□ 시간 부족\n□ 기타:")
            
            # 핵심 포인트
            point_cell = analysis_table.cell(0, 1)
            point_cell.width = Inches(3.25)
            point_p = point_cell.paragraphs[0]
            point_p.add_run("💡 핵심 포인트").bold = True
            point_p.add_run("\n\n")
            
            # 셀 높이 설정
            for cell in [cause_cell, point_cell]:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcH = OxmlElement('w:tcH')
                tcH.set(qn('w:val'), str(int(analysis_height * 1440)))
                tcH.set(qn('w:hRule'), 'exact')
                tcPr.append(tcH)
            
            # 문제 구분선 (마지막 문제 제외)
            if prob_num < problems_per_page - 1:
                separator = doc.add_paragraph("─" * 50)
                separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
                separator.paragraph_format.space_before = Pt(8)
                separator.paragraph_format.space_after = Pt(8)

# Streamlit 앱 설정
st.set_page_config(page_title="노트 양식 생성기", page_icon="📝", layout="wide")

st.title("📝 노트 양식 생성기")
st.markdown("다양한 노트 양식을 선택하고 Word 파일로 다운로드하세요!")

col1, col2 = st.columns([1, 2])

with col1:
    st.subheader("⚙️ 설정")
    
    # 노트 종류 선택
    notebook_type = st.selectbox(
        "노트 종류 선택",
        ["줄공책", "칸공책", "영어노트 (4선)", "코넬노트", "음악 오선지", 
         "한자노트", "다이어리", "달력", "수학 오답노트"]
    )
    
    # 사용자 정보 입력
    st.subheader("👤 사용자 정보")
    include_info = st.checkbox("사용자 정보 포함", value=True)
    
    if include_info:
        col_info1, col_info2 = st.columns(2)
        with col_info1:
            school_name = st.text_input("학교명", placeholder="예: 경북초등학교")
            student_name = st.text_input("이름", placeholder="예: 홍길동")
        with col_info2:
            grade = st.text_input("학년", placeholder="예: 3학년")
            class_num = st.text_input("반", placeholder="예: 2반")
    
    # 페이지 수
    if notebook_type not in ["다이어리", "달력"]:
        num_pages = st.number_input("페이지 수", min_value=1, max_value=50, value=5)
    
    # 용지 방향
    orientation = st.radio("용지 방향", ["세로", "가로"])
    
    # 노트별 추가 설정
    if notebook_type == "줄공책":
        lines_per_page = st.slider("페이지당 줄 수", 10, 35, 25)
    elif notebook_type == "칸공책":
        rows = st.slider("행 수", 5, 25, 15)
        cols = st.slider("열 수", 5, 25, 15)
        st.info("💡 팁: 많은 칸을 만들면 생성 시간이 길어질 수 있습니다.")
    elif notebook_type == "영어노트 (4선)":
        lines_per_page = st.slider("페이지당 줄 수", 5, 15, 10)
    elif notebook_type == "음악 오선지":
        staves_per_page = st.slider("페이지당 오선 수", 8, 14, 12)
    elif notebook_type == "한자노트":
        rows_per_page = st.slider("페이지당 행 수", 5, 10, 8)
        chars_per_row = st.slider("행당 칸 수", 8, 12, 10)
    elif notebook_type == "다이어리":
        start_date = st.date_input("시작 날짜", datetime.now())
        num_days = st.number_input("일수", min_value=1, max_value=365, value=7)
    elif notebook_type == "달력":
        col_cal1, col_cal2 = st.columns(2)
        with col_cal1:
            year = st.number_input("연도", min_value=2020, max_value=2030, value=datetime.now().year)
        with col_cal2:
            month = st.number_input("시작 월", min_value=1, max_value=12, value=datetime.now().month)
        num_months = st.number_input("개월 수", min_value=1, max_value=12, value=12)
    elif notebook_type == "수학 오답노트":
        problems_per_page = st.slider("페이지당 문제 수", 1, 4, 3)
    
    # 생성 버튼
    if st.button("📄 노트 생성", use_container_width=True, type="primary"):
        with st.spinner("노트를 생성하고 있습니다..."):
            try:
                # Document 생성
                doc = Document()
                
                # 용지 방향 설정
                section = doc.sections[0]
                if orientation == "가로":
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width, section.page_height = section.page_height, section.page_width
                
                # 여백 설정
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
                
                # 사용자 정보 준비
                user_info = None
                if include_info:
                    user_info = {
                        "school_name": school_name,
                        "grade": grade,
                        "class_num": class_num,
                        "student_name": student_name
                    }
                
                # 선택된 노트 종류에 따라 생성
                if notebook_type == "줄공책":
                    create_lined_notebook(doc, lines_per_page, num_pages, user_info)
                elif notebook_type == "칸공책":
                    create_grid_notebook(doc, rows, cols, num_pages, user_info)
                elif notebook_type == "영어노트 (4선)":
                    create_english_notebook(doc, lines_per_page, num_pages, user_info)
                elif notebook_type == "코넬노트":
                    create_cornell_notebook(doc, num_pages, user_info)
                elif notebook_type == "음악 오선지":
                    create_music_staff(doc, staves_per_page, num_pages, user_info)
                elif notebook_type == "한자노트":
                    create_chinese_notebook(doc, rows_per_page, chars_per_row, num_pages, user_info)
                elif notebook_type == "다이어리":
                    create_diary(doc, start_date, num_days, user_info)
                elif notebook_type == "달력":
                    create_calendar(doc, year, month, num_months, user_info)
                elif notebook_type == "수학 오답노트":
                    create_math_error_notebook(doc, problems_per_page, num_pages, user_info)
                
                # 모든 페이지에 푸터 추가
                add_footer(doc)
                
                # 메모리에 저장
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                # 파일명 생성
                if notebook_type == "다이어리":
                    filename = f"{notebook_type}_{start_date.strftime('%Y%m%d')}_{num_days}일.docx"
                elif notebook_type == "달력":
                    filename = f"{notebook_type}_{year}년_{month}월_{num_months}개월.docx"
                else:
                    filename = f"{notebook_type}_{num_pages}페이지.docx"
                
                # 다운로드 버튼
                st.download_button(
                    label="📥 Word 파일 다운로드",
                    data=doc_io.getvalue(),
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
                st.success("✅ 노트가 성공적으로 생성되었습니다!")
                
            except Exception as e:
                st.error(f"❌ 오류가 발생했습니다: {str(e)}")
                st.info("다른 설정으로 다시 시도해보세요.")

with col2:
    st.subheader("📖 사용 방법")
    st.markdown("""
    1. **노트 종류 선택**: 원하는 노트 양식을 선택하세요.
    2. **페이지 수 설정**: 생성할 페이지 수를 입력하세요.
    3. **용지 방향 선택**: 세로 또는 가로 방향을 선택하세요.
    4. **추가 설정**: 노트 종류에 따라 줄 수, 칸 수 등을 조정하세요.
    5. **노트 생성**: '노트 생성' 버튼을 클릭하세요.
    6. **다운로드**: 생성된 Word 파일을 다운로드하세요.
    """)
    
    st.subheader("📝 노트 종류 설명")
    with st.expander("줄공책"):
        st.markdown("""
        - 일반적인 줄이 그어진 노트
        - 글쓰기, 일기, 메모 등에 적합
        - 페이지당 줄 수 조정 가능
        - 회색 선으로 구성
        """)
    
    with st.expander("칸공책"):
        st.markdown("""
        - 격자 모양의 칸으로 구성된 노트
        - 수학, 도표, 그래프 그리기에 적합
        - 행과 열의 수 조정 가능
        - 정사각형에 가까운 칸으로 구성
        """)
    
    with st.expander("영어노트 (4선)"):
        st.markdown("""
        - 영어 필기체 연습용 4선 노트
        - 알파벳 쓰기 연습에 최적화
        - 점선과 실선으로 구성
        - 기준선이 굵게 표시됨
        """)
    
    with st.expander("코넬노트"):
        st.markdown("""
        - 효과적인 학습을 위한 노트 양식
        - 핵심어/질문, 노트, 요약 영역으로 구분
        - 체계적인 학습 정리에 적합
        - 복습과 정리가 용이한 구조
        """)
    
    with st.expander("음악 오선지"):
        st.markdown("""
        - 음악 작곡과 악보 작성용 오선지
        - 페이지당 8~14개의 오선 배치 가능
        - 음표, 쉼표, 음악 기호 작성에 적합
        - 작곡, 편곡, 음악 수업에 활용
        """)
    
    with st.expander("한자노트"):
        st.markdown("""
        - 한자/한문 연습용 격자 노트
        - 각 칸에 십자 가이드라인 포함
        - 한자의 획순과 균형 연습에 최적화
        - 중국어, 일본어 문자 연습에도 활용 가능
        """)
    
    with st.expander("다이어리"):
        st.markdown("""
        - 일일 다이어리 양식
        - 날씨, 기분, 중요도 표시 영역
        - 시간대별 일정 관리
        - 일기 작성 공간
        - 감사 일기 섹션 포함
        - 원하는 날짜부터 시작 가능
        """)
    
    with st.expander("달력"):
        st.markdown("""
        - 월별 달력 양식
        - 각 날짜별 메모 공간 포함
        - 주말은 색상으로 구분 (토요일: 파란색, 일요일: 빨간색)
        - 하단에 월별 메모 공간
        - 연간 계획이나 월간 일정 관리에 활용
        """)
    
    with st.expander("수학 오답노트"):
        st.markdown("""
        - **체계적인 수학 문제 정리를 위한 전문 노트**
        - 문제 정보: 문제 번호, 날짜, 출처, 난이도 기록
        - 문제 작성 공간 (연한 파란색 배경)
        - 풀이 과정용 격자 노트 (계산 과정 정리에 최적화)
        - 오답 원인 체크리스트
        - 핵심 포인트 정리 공간
        - 페이지당 1~4문제 설정 가능
        - **활용 팁:**
          - 시험 후 틀린 문제를 체계적으로 정리
          - 오답 원인을 분석하여 같은 실수 방지
          - 핵심 개념과 풀이법을 한눈에 정리
          - 시험 전 복습 자료로 활용
        """)
    
    st.subheader("💡 추가 팁")
    st.markdown("""
    - **인쇄 시**: 프린터 설정에서 '실제 크기'로 인쇄하세요.
    - **양면 인쇄**: 용지 절약을 위해 양면 인쇄를 권장합니다.
    - **PDF 변환**: Word 파일을 PDF로 변환하면 레이아웃이 더 안정적입니다.
    - **문제 해결**: 생성이 안 되면 페이지 수나 칸 수를 줄여보세요.
    - **수학 오답노트 활용법**: 
      - 문제를 풀 때 사용한 개념과 공식을 함께 정리하세요
      - 비슷한 유형의 문제를 함께 모아서 정리하면 효과적입니다
      - 정기적으로 복습하여 실수를 줄여나가세요
    """)
